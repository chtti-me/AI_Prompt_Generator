# -*- coding: utf-8 -*-
"""
MD 轉 Word 腳本 - 將 Markdown 文件轉換為 Word 並嵌入截圖
"""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 路徑設定
MD_FILE = r"D:\自己架設AI_零基礎到大師\CH5_AI Prompt互動提示詞生成系統\book\CH5_AI Prompt互動提示詞生成系統.md"
TEMPLATE_FILE = r"D:\自己架設AI_零基礎到大師\書籍範本\reference.docx"
OUTPUT_FILE = r"D:\自己架設AI_零基礎到大師\CH5_AI Prompt互動提示詞生成系統\book\CH5_AI Prompt互動提示詞生成系統.docx"
SCREENSHOT_DIR = r"D:\自己架設AI_零基礎到大師\CH5_AI Prompt互動提示詞生成系統\截圖"

# 截圖對應表
SCREENSHOT_MAP = {
    'ch5_startup_success.png': 'ch5_startup_success.png',
    'ch5_env_config.png': 'ch5_env_config.png',
    'ch5_homepage.png': 'ch5_homepage.png .png',  # 注意檔名有空格
    'ch5_generator_form.png': 'ch5_generator_form.png',
    'ch5_generated_result.png': 'ch5_generated_result.png',
    'ch5_rag_upload.png': 'ch5_rag_upload.png',
    'ch5_rag_search.png': 'ch5_rag_search.png',
    'ch5_dashboard.png': 'ch5_dashboard.png',
    'ch5_optimize_result.png': 'ch5_optimize_result.png',
    'ch5_rag_search_detail.png': 'ch5_rag_search_detail.png',
    'ch5_service_account.png': 'ch5_service_account.png',
    'ch5_sheet_sharing.png': 'ch5_sheet_sharing.png',
    'ch5_gas_menu.png': 'ch5_gas_menu.png',
    'ch5_sheets_collaboration.png': 'ch5_sheets_collaboration.png',
    'ch5_memory_usage.png': 'ch5_memory_usage.png',
    'ch5_final_dashboard.png': 'ch5_final_dashboard.png',
}


def set_chinese_font(run, font_name='微軟正黑體', font_size=12):
    """設定中文字體"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def add_heading(doc, text, level):
    """新增標題（不使用內建樣式）"""
    para = doc.add_paragraph()
    run = para.add_run(text)

    # 根據層級設定字體大小
    if level == 1:
        font_size = 18
    elif level == 2:
        font_size = 16
    elif level == 3:
        font_size = 14
    else:
        font_size = 12

    set_chinese_font(run, '微軟正黑體', font_size)
    run.bold = True

    # 設定段落間距
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)

    return para


def add_paragraph_text(doc, text):
    """新增段落文字"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run, '微軟正黑體', 12)
    return para


def add_code_block(doc, code, language=''):
    """新增程式碼區塊"""
    para = doc.add_paragraph()
    para.style = 'Normal'
    # 設定灰色背景效果（透過縮排模擬）
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.right_indent = Cm(0.5)

    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return para


def add_image(doc, image_path, caption=''):
    """新增圖片"""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(5.5))

        # 新增圖片說明
        if caption:
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(caption)
            set_chinese_font(caption_run, '微軟正黑體', 10)
            caption_run.italic = True

        print(f"  [OK] 插入圖片: {os.path.basename(image_path)}")
        return True
    else:
        print(f"  [警告] 找不到圖片: {image_path}")
        return False


def find_screenshot(filename):
    """尋找截圖檔案"""
    # 先嘗試直接對應
    if filename in SCREENSHOT_MAP:
        actual_filename = SCREENSHOT_MAP[filename]
        path = os.path.join(SCREENSHOT_DIR, actual_filename)
        if os.path.exists(path):
            return path

    # 嘗試直接路徑
    path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(path):
        return path

    # 模糊搜尋
    base_name = filename.replace('.png', '')
    for f in os.listdir(SCREENSHOT_DIR):
        if base_name in f:
            return os.path.join(SCREENSHOT_DIR, f)

    return None


def parse_and_convert(md_content, doc):
    """解析 Markdown 並轉換為 Word"""
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    code_language = ''
    image_count = 0

    while i < len(lines):
        line = lines[i]

        # 處理程式碼區塊
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_language = line[3:].strip()
                code_buffer = []
            else:
                in_code_block = False
                if code_buffer:
                    add_code_block(doc, '\n'.join(code_buffer), code_language)
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # 處理標題
        if line.startswith('# '):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:].strip(), 4)

        # 處理截圖註記
        elif '【註記】請截圖' in line:
            # 提取截圖檔名
            match = re.search(r'`([^`]+\.png)`', line)
            if match:
                filename = match.group(1)
                image_path = find_screenshot(filename)

                # 提取說明文字
                desc_match = re.search(r'【註記】請截圖「([^」]+)」', line)
                caption = desc_match.group(1) if desc_match else ''

                if image_path:
                    image_count += 1
                    add_image(doc, image_path, f"圖 5-{image_count}：{caption}")
                else:
                    # 插入佔位符
                    para = add_paragraph_text(doc, f"[圖片：{caption} - {filename}]")
                    para.runs[0].italic = True

        # 處理分隔線
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50)

        # 處理引用區塊
        elif line.startswith('> '):
            text = line[2:].strip()
            if '【註記】' not in text:  # 跳過截圖註記的引用
                para = add_paragraph_text(doc, text)
                para.paragraph_format.left_indent = Cm(1)

        # 處理一般段落
        elif line.strip():
            # 處理粗體
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', line)
            # 處理行內程式碼
            text = re.sub(r'`([^`]+)`', r'「\1」', text)
            add_paragraph_text(doc, text)

        # 空行
        else:
            pass  # 跳過空行

        i += 1

    return image_count


def main():
    print("=" * 60)
    print("MD 轉 Word 轉換工具")
    print("=" * 60)

    # 讀取 Markdown 檔案
    print(f"\n[1] 讀取 Markdown 檔案...")
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()
    print(f"    檔案大小: {len(md_content)} 字元")

    # 載入範本或建立新文件
    print(f"\n[2] 載入 Word 範本...")
    if os.path.exists(TEMPLATE_FILE):
        doc = Document(TEMPLATE_FILE)
        # 清除範本內容，保留樣式
        for element in doc.element.body[:]:
            doc.element.body.remove(element)
        print(f"    使用範本: {TEMPLATE_FILE}")
    else:
        doc = Document()
        print("    [警告] 範本不存在，使用預設格式")

    # 轉換內容
    print(f"\n[3] 轉換 Markdown 為 Word...")
    image_count = parse_and_convert(md_content, doc)
    print(f"    共插入 {image_count} 張圖片")

    # 儲存文件
    print(f"\n[4] 儲存 Word 文件...")
    doc.save(OUTPUT_FILE)
    print(f"    輸出檔案: {OUTPUT_FILE}")

    print("\n" + "=" * 60)
    print("轉換完成！")
    print("=" * 60)


if __name__ == '__main__':
    main()
